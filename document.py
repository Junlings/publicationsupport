from contents.paragraph import Paragraph, ParagraphLib, SingleSentence
from contents.figures import FigureLib
from contents.table import TableLib
from contents.equations import EqLib
from contents.documenrstruct import DocumentStruct
from contents.latex_meta_lib import metacls_objlib
from uti.extract import extract_labels
from uti.latexmath2png import math2png, table2png
from uti.filefolder import ensure_dir, copyfile
import re
from docx import docx #contenttypes, savedocx, relationshiplist, newdocument, nsprefixes, paragraph, coreproperties, appproperties, websettings
import os
import glob
import Image
import cPickle as pickle
import random, string


def randomword(N):
   return ''.join(random.choice(string.ascii_uppercase + string.digits) for x in range(N))




class line():
    def __init__(self,seq=-1,content='',tag='',ltype='N/A'):
        self.seq = seq
        self.content = content
        self.tag = tag
        self.type = ltype
        self.lib = {}
    
    def extract(self,label,quote=['{','}']):
        """
        Extract the parameter after the label{}
        """
        
        pattern1 = re.compile(label)

class LineLib():
    __metaclass__ = metacls_objlib
    
    ''' variable library'''
    def __init__(self):
        #self.tag = tag
        self.itemlib = {}
        self.count = 0
        self.libtype = line    
    
    def __str__(self):
        output = ''
        for key,line in self.itemlib.items():
            
            output += '[' + str(line.seq) + ']'+'[' + str(line.type) + ']'
            output += line.content
            output += '\n'
        return  output
    
class document():
    
    def __init__(self):
        self.struct = DocumentStruct()
        self.figurelib = FigureLib()
        self.tablelib = TableLib()
        self.eqlib = EqLib()
        self.LineLib = LineLib()
        self.paralib = ParagraphLib()
        self.tag = 'default'
    
    def GetTreeDict(self):
        ''' obtain the document tree'''
        treedict = {}
        
        treedict = self.figurelib.GetLibDict('figure',treedict)
        treedict = self.tablelib.GetLibDict('table',treedict)
        treedict = self.eqlib.GetLibDict('equation',treedict)
        treedict = self.paralib.GetLibDict('paragraph',treedict)
        
        return treedict
   
    def import_tex(self,filename):
        """
        import tex based files
        """
        self.workdir = os.path.dirname(filename)
        
        
        if '.' not in filename:
            filename=filename+'.tex'
        f_full=open(filename,'rU')
        lines = f_full.readlines()
        
        

        nline = 1
        for i in range(0,len(lines)):
            tempstr = lines[i]
            tempstr.strip()
            if tempstr.endswith('\n'):
                tempstr = tempstr[:len(tempstr)-1]
                
            if len(tempstr) > 1:
                
                if '\\input' not in tempstr:
                    templine = line(tag=nline,seq=nline,content=tempstr)
                    self.LineLib.Add(templine) # this function is defined in the meta class defination
                    nline += 1
                    
                else:  # get extended file from inputs command
                    filepath = tempstr[tempstr.index('{')+1:tempstr.index('}')]
                    tempdoc = document()
                    tempdoc.import_tex(os.path.join(self.workdir,filepath))
                    
                    for key,item in tempdoc.LineLib.itemlib.items():
                        templine = line(tag=nline,seq=nline,content=item.content)
                        if len(item.content) <5:
                            pass
                        else:
                            self.LineLib.Add(templine) # this function is defined in the meta class defination
                            nline += 1                            

                    
            preline = tempstr

        
    def labeling(self):
        ''' judge the line type by the typical words '''
        nline = 0
        while 1:
            try:
                nline += 1
                line = self.LineLib.itemlib[nline]
            except:
                break
            else:
                if line.content.startswith(r'%'):
                    line.type = 'comments'
                elif line.content.startswith('\\input'):
                    line.type = 'input'
                elif line.content.startswith('\\include'):
                    line.type = 'include'
                elif line.content.startswith('\\'):
                    if '{' in line.content:
                        keys = re.findall(r"\\(.*){(.*)}",line.content)
                        if keys != []:
                            line.type = 'label:%s:%s' % keys[0]
                        else:
                            line.type = 'undefined'
                    else:
                        keys = re.findall(r"\\(.*)",line.content)
                        line.type = 'label:%s' % keys[0]
                else:
                    line.type = 'text'
    
    
    def AddParagraph(self,lines):
        pass
        
    
    def AddFigureFromTex(self,lines):
        f1 = self.figurelib.AddByTex(lines)
        return f1
    
    def AddEquationFromTex(self,lines):
        e1 = self.eqlib.AddByTex(lines)
        return e1
        
    def AddTableFromTex(self,lines):
        t1 = self.tablelib.AddByTex(lines)
        return t1

    def AddParagraphFromTex(self,tag,lines):
        p1 = self.paralib.AddByTex(tag,lines)
        return p1 
    
    def AddItemsByTex(self,current_paragraph,current_paragraph_node):
        
        itemlist = extract_labels(current_paragraph.wholetext)
        if itemlist != None:
            current_paragraph.AddByDict(itemlist)
        
    
    def extract(self):
        """
        Extract the content from the tex inputs
        """
        self.labeling()
        
        line_content = 'text' 
        tablecontentlist=[]
        figurecontentlist=[]
        equationcontentlist=[]
        maincontentlist=[]
        
        current_section = None
        current_subsection = None
        current_subsubsection = None
        current = None
        current_paragraph = None
        current_paragraph_node = None
        
        for key,line in self.LineLib.itemlib.items():
            #line = self.LineLib.itemlib[i]
            
            if line.content[0] == '%':  # comments
                continue
            
            if '\section' in line.content:
                text = line.content[line.content.index('{')+1:line.content.index('}')]
                tag = 'section'
                current_section = self.struct.AddSection(self.struct.root,tag,text)
                current = current_section
                line_content='section'
                
            elif '\subsection' in line.content:
                text = line.content[line.content.index('{')+1:line.content.index('}')]
                tag = 'subsection'
                current_subsection = self.struct.AddSection(current_section,tag,text)
                current = current_subsection
                line_content='section'

            elif '\subsubsection' in line.content:
                text = line.content[line.content.index('{')+1:line.content.index('}')]
                tag = 'subsubsection'
                current_subsubsection = self.struct.AddSection(current_subsection,tag,text)
                current = current_subsubsection
                line_content='section'
            
            elif  '\\begin{table}' in line.content:
                line_content='table'
                temp_table=[]
                
            elif  '\\end{table}' in line.content:
                line_content='tabletomain'
                
                
            elif  '\\begin{figure}' in line.content:
                line_content='figure'
                temp_figure=[]
                
            elif  '\\end{figure}' in line.content:
                line_content='figuretomain'
                
                
            elif  '\\begin{equation}' in line.content:
                line_content='equation'
                temp_equation=[]
                
            elif  '\\end{equation}' in line.content:
                line_content='equationtomain'
            
            elif line_content == 'section':
                line_content='text'
            
            #elif  line.content[0]=='\\':
            #    line_content='other'

            else:
                pass
               

            
            line.type = line_content
            ## apply action for each category
            if line_content=='table':
                temp_table.append(line)
                
            elif line_content=='tabletomain':
                temp_table.append(line)
                tempobj = self.AddTableFromTex(temp_table)
                tablecontentlist.append(temp_table)
                line_content='text'
                
            elif line_content=='figure':
                temp_figure.append(line)
            
            elif line_content=='figuretomain':
                temp_figure.append(line)
                tempobj = self.AddFigureFromTex(temp_figure)#figure({'content':temp_figure,'seq':temp_figure[0].seq})
                line_content='text'
                figurecontentlist.append(temp_figure)
            
            elif line_content=='equation':
                temp_equation.append(line)
            
            elif line_content=='equationtomain':
                temp_equation.append(line)
                self.AddEquationFromTex(temp_equation)
                equationcontentlist.append(temp_equation)
                line_content='text'
                
            elif line_content == 'text' and line.content[0] != '\\':
                #tag = '_'.join(line.content.split(' ')[0:5])
                tag = 'PARA'+randomword(10)
                #print tag
                current_paragraph = self.AddParagraphFromTex(tag,line.content)
                if current != None:  # not support abstract for now
                    current_paragraph_node = self.struct.AddParagraph(current,current_paragraph)
                    self.AddItemsByTex(current_paragraph,current_paragraph_node)
            else:
                pass

        return 1
    
    
    def GetPara(self,label):
        if label in self.paralib.itemlib.keys():
            return self.paralib.itemlib[label]
        else:
            return None
    
    
    
    def export_struc(self,filename):
        ''' export the xml document structure'''
        self.struct.exportfile(filename)

    
    
    def docx_export(self,dirname,filename):
        self.docx_new()
        self.docx_add_by_tree()
        self.docx_add_figure_by_list()
        self.docx_add_table_by_list()
        self.docx_add_equation_by_list()
        self.docx_save(os.path.join(dirname, filename))        
    
    def docx_new(self):
        
    
        # Make a new document tree - this is the main part of a Word document
        mydocument = docx.newdocument()
    
        # This xpath location is where most interesting content lives
        body = mydocument.xpath('/w:document/w:body', namespaces=docx.nsprefixes)[0]             
        relationships = docx.relationshiplist()
        
        self.docx = {'document':mydocument,'body':body,'relationships':relationships,
                     'figurelist':[],'tablelist':[],'equationlist':[],'referencelist':[]
                     }
    
    
    def docx_add_by_tree(self):
        for node in self.struct.root.xpath("//*"):
            if node.tag == 'section':
                self.docx_add_section(node.text,1)
            elif node.tag == 'subsection':
                self.docx_add_section(node.text,2)
            elif node.tag == 'subsubsection':
                self.docx_add_section(node.text,3)
            elif node.tag == 'paragraph':
                para = self.GetPara(node.text)
                self.docx_add_paragraph('['+para.tag+']'+para.wholetext)
                
                # record figure list
                self.docx['figurelist'].extend(para.itemdict['Figure'])
                # record figure list
                self.docx['tablelist'].extend(para.itemdict['Table'])
                # record figure list
                self.docx['equationlist'].extend(para.itemdict['Equation'])
                # record figure list
                self.docx['referencelist'].extend(para.itemdict['Reference'])
            else:
                pass
     
    def docx_add_section(self,sectiontext,level):
        self.docx['body'].append(docx.heading(sectiontext, level))
        
    def docx_add_paragraph(self,paragraphtext):
        self.docx['body'].append(docx.paragraph(paragraphtext))
    
    
    def docx_add_figure(self,imagefilename,description):
        self.docx['relationships'], picpara = docx.picture(self.docx['relationships'], imagefilename,
                                         description)
        self.docx['body'].append(docx.pagebreak(type='page', orient='portrait'))
        self.docx['body'].append(picpara)
        self.docx['body'].append(docx.paragraph(description))
        
    def docx_add_figure_from_lib(self,key):
        figure = self.figurelib.itemlib[key]
        
        if 'eps' in figure.itemlib.keys():
            singlefigure = figure.itemlib['eps']
            imgfile = os.path.join(self.exportdst,self.exportfolder['figure'],singlefigure.path)
            self.docx_add_figure(imgfile,figure.caption)
        

    def docx_add_eq_from_lib(self,key):
        equation = self.eqlib.itemlib[key]
        
        if 'png' in equation.itemlib.keys():
            singleeq = equation.itemlib['png']
            imgfile = os.path.join(self.exportdst,self.exportfolder['equation'],singleeq.path)
            self.docx_add_figure(imgfile,equation.caption)
        elif 'tex' in equation.itemlib.keys():
            self.docx_add_paragraph(equation.caption)  
            self.docx_add_paragraph(equation.itemlib['tex'].latex)            
            
    def docx_add_table_from_lib(self,key):
        table = self.tablelib.itemlib[key]
        
        if 'png' in table.itemlib.keys():
            singletable = table.itemlib['png']
            imgfile = os.path.join(self.exportdst,self.exportfolder['table'],singletable.path)
            #self.docx_add_paragraph(table.caption)  
            self.docx_add_figure(imgfile,table.caption)
            
        elif 'tex' in table.itemlib.keys():
            self.docx_add_paragraph(table.itemlib['tex'].latex)           
    
    def docx_add_figure_by_list(self):
        
        for key in self.docx['figurelist']:
            self.docx_add_figure_from_lib(key)
        
    def docx_add_table_by_list(self):
        
        for key in self.docx['tablelist']:
            self.docx_add_table_from_lib(key)        

    def docx_add_equation_by_list(self):
        
        for key in self.docx['equationlist']:
            self.docx_add_eq_from_lib(key)        
    
        
    def docx_cofigure(self):
        pass
    
    def docx_save(self,filename,title='title',subject='subject',creator='unknown',keywords=[]):
   
        
        title    = title #'Python docx demo'
        subject  = subject #'A practical example of making docx from Python'
        creator  = creator #'Mike MacCana'
        keywords = keywords #['python', 'Office Open XML', 'Word']
        
        coreprops = docx.coreproperties(title=title, subject=subject, creator=creator,
                                   keywords=keywords)
        appprops = docx.appproperties()
        contenttypes = docx.contenttypes()
        websettings = docx.websettings()
        
        
        
        wordrelationships = docx.wordrelationships(self.docx['relationships'])
        
    
        # Save our document
        docx.savedocx(self.docx['document'], coreprops, appprops, contenttypes, websettings,
                 wordrelationships, filename)
    
    
    
    def EqTex2Png(self):
        
        for key,item in self.eqlib.itemlib.items():
            latextext = '\n'.join(item.itemlib['tex'].latex)
            latextext = '\[' + latextext + '.\]'
            #print latextext
            math2png([latextext],os.path.join(self.exportdst,self.exportfolder['equation'],item.tag),prefix=item.tag)
            item.AddPng(os.path.join(item.tag,item.tag+'.png'))
    
    def TableTex2Png(self):
        for key,item in self.tablelib.itemlib.items():
            latextext = item.itemlib['tex'].latex
            #latextext = '\[' + latextext + '.\]'
            #print latextext
            tablename = os.path.join(self.exportdst,self.exportfolder['table'],item.tag)
            table2png([latextext],tablename,prefix=item.tag)        
            
            # add png file format to the library
            item.AddPng(os.path.join(item.tag,item.tag+'.png'))
            
    def FigureEps2Png(self):
        for key,item in self.figurelib.itemlib.items():
            epsfile = os.path.join(self.exportdst,self.exportfolder['figure'],item.tag,item.tag + '.eps')
            pngfile = os.path.join(self.exportdst,self.exportfolder['figure'],item.tag,item.tag + '.png')
            
            try:
                im = Image.open(epsfile)
                im.save(pngfile)
            except:
                pass
            
    
    
    def FigureImportFolder(self,folder):
        ''' import image from desitnation folder'''
        for root, subdirs, files in os.walk(folder):
            for file in files:
                if os.path.splitext(file)[1].lower() in ('.jpg', '.jpeg','.JPG','.png','.PNG','.eps','.EPS'):
                     print os.path.join(root, file)
                     

    
    def ExportProject(self,desinationfolder):
        ''' save project in destination folder'''
        self.exportdst = desinationfolder
        self.exportfolder = {'equation':'equation','figure':'img','table':'table','paragraph':'paragraph'}
        
        # export all components
        self.ExportParaLib(os.path.join(desinationfolder,self.exportfolder['paragraph']))
        self.ExportFigureLib(os.path.join(desinationfolder,self.exportfolder['figure']))
        self.ExportTableLib(os.path.join(desinationfolder,self.exportfolder['table']))
        self.ExportEqLib(os.path.join(desinationfolder,self.exportfolder['equation']))
        
        filename = os.path.join(self.exportdst,self.tag+'.xml')
        self.export_struc(filename)
        
        fp = open(os.path.join(self.exportdst,self.tag+'.lib'),'w')
        
        lib = {'equation':self.eqlib,'figure':self.figurelib,'table':self.tablelib,'paragraph':self.paralib}
        pickle.dump(lib,fp)
        
        

    def ExportParaLib(self,parafolder):
        ensure_dir(parafolder)
        
        for key, item in self.paralib.itemlib.items():
            #dstroot = os.path.join(parafolder,key)
            #ensure_dir(dstroot)
            
            fp = open(os.path.join(parafolder,item.tag+'.tex'),'w')
            
            fp.write(item.wholetext)
            fp.close()
            
    
    def ExportFigureLib(self,figurefolder):
        ''' save figure to figure folder  '''
        ensure_dir(figurefolder)
        
        for key, item in self.figurelib.itemlib.items():
            dstroot = os.path.join(figurefolder,key)
            ensure_dir(dstroot)
            
            for singlefigurekey,singlefigure in item.itemlib.items():
                scrfile = os.path.join(self.workdir,singlefigure.path)
                filename = key +'.' + singlefigurekey
                
                copyfile(scrfile,filename,dstroot)
                singlefigure.path = os.path.join(key,filename)
            
    def ExportTableLib(self,tablefolder):
        ensure_dir(tablefolder)
        
        for key, item in self.tablelib.itemlib.items():
            dstroot = os.path.join(tablefolder,key)
            ensure_dir(dstroot)
            
            for singletablekey,singletable in item.itemlib.items():
                
                if singletable.filepath != None:
                    ''' direct copy table '''
                    scrfile = os.path.join(self.workdir,singletablekey,singletable.path)
                    filename = key +'.' + singletablekey
                    copyfile(scrfile,filename,dstroot)
                    
                else:
                    ''' write table file '''
                    pass
            filename = key +'.' + 'tex'
            self.TexWriteTableFile(key,os.path.join(tablefolder,key,filename))
            singletable.path = os.path.join(key,filename) 
        
    def TexWriteTableFile(self,tablelabel,filename):
        ''' write table to latex file '''
        tableobj = self.tablelib.itemlib[tablelabel]
        
        fd = open(filename,'w')
        temp = '\\begin{table}\n\caption{%s}\n\label{%s}\centering\n' % (tableobj.caption,tableobj.tag)
        fd.write(temp)
        fd.write(tableobj.itemlib['tex'].latex)
        #for line in tableobj.itemlib['tex'].latex:
        #    fd.write(line)
        fd.write('\n')
        fd.write('\end{table}\n')   
        
    def ExportEqLib(self,eqfolder):
        ensure_dir(eqfolder)
        
        for key, item in self.eqlib.itemlib.items():
            dstroot = os.path.join(eqfolder,key)
            ensure_dir(dstroot)
            
            for singleeqkey,singleeq in item.itemlib.items():
                
                if singleeq.path != None:
                    ''' direct copy table '''
                    scrfile = os.path.join(self.workdir,singleeqkey,singleeq.path)
                    filename = key +'.' + singleeqkey
                    copyfile(scrfile,filename,dstroot)
                    
                else:
                    ''' write table file '''
                    pass
            filename = key +'.' + 'tex'
            self.TexWriteEqFile(key,os.path.join(eqfolder,key,filename))        

    def TexWriteEqFile(self,eqlabel,filename):
        ''' write table to latex file '''
        eqobj = self.eqlib.itemlib[eqlabel]
        
        fd = open(filename,'w')
        
        temp = '\\begin{equation}\label{%s}\n' % (eqobj.tag)
        fd.write(temp)
        #fd.write(tableobj.itemlib['tex'].latex)
        for line in eqobj.itemlib['tex'].latex:
            fd.write(line)
        fd.write('\n')
        fd.write('\end{equation} \n')
        
if __name__ == '__main__':
    
    d1 = document()
    d1.import_tex('test/manuscript.tex')
    
    
    d1.extract()
    print d1.LineLib
    aa = d1.struct.pretty()
    
    f = open('test/temp.xml','w')
    f.write(aa)
    f.close()
    
    d1.docx_new()
    
    
    for key, item in d1.paralib.itemlib.items():
        try:
            print d1.paralib.itemlib[key].itemdict
        except:
            pass
        
    
    
    d1.ExportProject(r'M:\github\publicationsupport\test\export')
    
    d1.TableTex2Png()
    d1.docx_add_by_tree()
    d1.docx_add_figure_by_list()
    d1.docx_add_table_by_list()
    #d1.docx_add_figure(r'M:\github\publicationsupport\test\export\img\Fig_compout_setup4\Fig_compout_setup4.png','this is a test')
    #d1.docx_save(r'M:\github\publicationsupport\test\export\text.docx')
    
    d1.EqTex2Png()
    #d1.FigureImportFolder(r'M:\github\publicationsupport\test\img')
    
    
    #d1.EqTex2Png()
    #d1.TableTex2Png()
    #d1.FigureEps2Png()
    print 1